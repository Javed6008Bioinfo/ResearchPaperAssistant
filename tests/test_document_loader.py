"""
Test suite for the document_loader service.

Covers: TXT loading (with encoding fallback), DOCX loading (paragraphs +
tables), PDF loading (structure/page handling), unsupported file types,
and file storage collision-safe naming.

Sample files are generated programmatically in a temp directory rather
than committed as binary fixtures, keeping the repo clean and tests
fully self-contained.
"""

from pathlib import Path

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from app.services.document_loader import load_document, save_uploaded_file
from app.utils.exceptions import DocumentParsingError, UnsupportedFileTypeError


# --- Fixtures: generate real sample files on the fly ---

@pytest.fixture
def sample_txt(tmp_path: Path) -> Path:
    file_path = tmp_path / "sample.txt"
    file_path.write_text("This is a sample research paper abstract.\n\nMethods section text.")
    return file_path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    file_path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("Abstract: This study investigates gene expression.")
    doc.add_paragraph("Methods: RNA-seq was performed on 20 samples.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Gene"
    table.rows[0].cells[1].text = "Expression"
    table.rows[1].cells[0].text = "BRCA1"
    table.rows[1].cells[1].text = "High"
    doc.save(str(file_path))
    return file_path


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """
    Creates a minimal valid (blank-page) PDF. pypdf can WRITE pages but
    cannot draw text without an additional rendering library, so this
    fixture validates structural handling (page counting, no crash on
    a legitimate PDF) rather than text-content extraction.
    """
    file_path = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    with open(file_path, "wb") as f:
        writer.write(f)
    return file_path


@pytest.fixture
def empty_txt(tmp_path: Path) -> Path:
    file_path = tmp_path / "empty.txt"
    file_path.write_bytes(b"")
    return file_path


@pytest.fixture
def unsupported_file(tmp_path: Path) -> Path:
    file_path = tmp_path / "sample.csv"
    file_path.write_text("gene,expression\nBRCA1,high")
    return file_path


# --- TXT loader tests ---

def test_load_txt_success(sample_txt: Path) -> None:
    document = load_document(sample_txt)
    assert document.file_type == "txt"
    assert "sample research paper abstract" in document.full_text
    assert document.page_count == 1
    assert not document.is_empty()


def test_load_txt_latin1_fallback(tmp_path: Path) -> None:
    file_path = tmp_path / "latin1.txt"
    # Write bytes that are invalid UTF-8 but valid latin-1
    file_path.write_bytes(b"Caf\xe9 gene expression results")
    document = load_document(file_path)
    assert "gene expression results" in document.full_text


# --- DOCX loader tests ---

def test_load_docx_success(sample_docx: Path) -> None:
    document = load_document(sample_docx)
    assert document.file_type == "docx"
    assert "gene expression" in document.full_text.lower()
    assert "BRCA1" in document.full_text  # confirms table extraction worked
    assert document.page_count == 1


# --- PDF loader tests ---

def test_load_pdf_success(sample_pdf: Path) -> None:
    document = load_document(sample_pdf)
    assert document.file_type == "pdf"
    assert document.page_count == 2  # two blank pages created in fixture
    assert document.is_empty()  # blank pages have no extractable text


# --- Error handling tests ---

def test_load_empty_file_raises(empty_txt: Path) -> None:
    with pytest.raises(DocumentParsingError):
        load_document(empty_txt)


def test_load_unsupported_type_raises(unsupported_file: Path) -> None:
    with pytest.raises(UnsupportedFileTypeError):
        load_document(unsupported_file)


def test_load_nonexistent_file_raises(tmp_path: Path) -> None:
    missing_path = tmp_path / "does_not_exist.pdf"
    with pytest.raises(DocumentParsingError):
        load_document(missing_path)


# --- File storage tests ---

def test_save_uploaded_file_creates_unique_names() -> None:
    path_1 = save_uploaded_file(b"content one", "paper.txt")
    path_2 = save_uploaded_file(b"content two", "paper.txt")

    assert path_1.exists()
    assert path_2.exists()
    assert path_1 != path_2  # collision-safe naming confirmed
    assert path_1.read_bytes() == b"content one"
    assert path_2.read_bytes() == b"content two"

    # cleanup
    path_1.unlink()
    path_2.unlink()


def test_save_uploaded_file_rejects_unsupported_type() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        save_uploaded_file(b"data", "malicious.exe")