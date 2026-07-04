"""
DOCX document loader using python-docx.

DOCX has no fixed "page" concept in the underlying XML (pagination is a
rendering-time detail in Word), so we treat the whole document as a
single logical page for the `pages` field, while still extracting
paragraph-level text faithfully and preserving table content.
"""

from pathlib import Path

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.services.document_loader.base_loader import BaseDocumentLoader
from app.services.document_loader.schemas import LoadedDocument
from app.utils.exceptions import DocumentParsingError
from app.utils.logger import logger


class DOCXLoader(BaseDocumentLoader):
    """Extracts text (paragraphs + tables) from DOCX files."""

    def load(self, file_path: Path) -> LoadedDocument:
        try:
            docx_file = DocxDocument(str(file_path))
        except (PackageNotFoundError, OSError) as exc:
            raise DocumentParsingError(
                f"Failed to open DOCX '{file_path.name}': {exc}"
            ) from exc

        paragraph_texts = [p.text.strip() for p in docx_file.paragraphs if p.text.strip()]
        table_texts = self._extract_table_text(docx_file)

        full_text = "\n\n".join(paragraph_texts + table_texts).strip()

        core_props = docx_file.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
        }

        document = LoadedDocument(
            filename=file_path.name,
            file_type="docx",
            full_text=full_text,
            pages=[full_text],  # DOCX has no fixed pagination at the XML level
            page_count=1,
            metadata=metadata,
        )

        if document.is_empty():
            logger.warning(f"DOCX '{file_path.name}' produced no extractable text.")

        return document

    @staticmethod
    def _extract_table_text(docx_file: DocxDocument) -> list[str]:
        """
        Flattens table content into readable text blocks.

        Research papers often contain data tables (e.g., results, gene
        lists) — skipping tables would silently drop important content.
        """
        table_blocks: list[str] = []
        for table in docx_file.tables:
            rows_text = []
            for row in table.rows:
                cells_text = [cell.text.strip() for cell in row.cells]
                rows_text.append(" | ".join(cells_text))
            if rows_text:
                table_blocks.append("\n".join(rows_text))
        return table_blocks